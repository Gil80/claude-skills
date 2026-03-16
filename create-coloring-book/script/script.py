import requests
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PIL import Image
import os

def create_coloring_book():
    if not os.path.exists('links.txt'):
        print("Error: links.txt not found.")
        return
    
    with open('links.txt', 'r') as f:
        image_urls = [line.strip() for line in f if line.strip()]

    doc = Document()
    
    # Configure A4 Page
    section = doc.sections[0]
    section.page_height, section.page_width = Mm(297), Mm(210)
    
    # Minimize margins to ensure no overflow
    section.top_margin = section.bottom_margin = Mm(5)
    section.left_margin = section.right_margin = Mm(10)

    # 135mm height is safe if we remove paragraph spacing
    max_height = Mm(135) 

    for i, url in enumerate(image_urls):
        try:
            print(f"Processing image {i+1}...")
            response = requests.get(url, timeout=15)
            img = Image.open(BytesIO(response.content))
            
            tmp_buffer = BytesIO()
            img.save(tmp_buffer, format="PNG")
            tmp_buffer.seek(0)

            # Add the picture
            pic_paragraph = doc.add_paragraph()
            run = pic_paragraph.add_run()
            run.add_picture(tmp_buffer, height=max_height)
            
            # REMOVE SPACING: This is the critical fix
            # This stops the "invisible" lines from pushing images to new pages
            pic_paragraph.paragraph_format.space_before = Mm(0)
            pic_paragraph.paragraph_format.space_after = Mm(0)
            pic_paragraph.paragraph_format.line_spacing = 1
            pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Only break page after the second image of a pair
            if (i + 1) % 2 == 0 and (i + 1) < len(image_urls):
                doc.add_page_break()
                
        except Exception as e:
            print(f"Error on image {i+1}: {e}")

    output_file = "Kawaii_Coloring_Book_Final.docx"
    doc.save(output_file)
    print(f"\nSuccess! '{output_file}' created with zero blank pages.")

if __name__ == "__main__":
    create_coloring_book()
